"""Document text extraction for RAG (Part 21 / §37).

Supports PDF, TXT, CSV, JSON, DOCX. Optional parsers are imported lazily
so the core pipeline never hard-depends on them; a missing parser reports
what to install rather than silently returning empty text.

Page/section numbers are recorded ONLY when the parser actually reports
them. They are never inferred or invented (Part 21).
"""

import csv
import io
import json
import os
from dataclasses import dataclass, field


@dataclass
class DocumentSection:
    text: str
    # None when the source format has no real page concept.
    page: int = None
    section: str = None
    source_file: str = ""

    def citation(self) -> str:
        name = os.path.basename(self.source_file)
        if self.page is not None:
            return f"{name}, page {self.page}"
        if self.section:
            return f"{name}, {self.section}"
        return name


@dataclass
class LoadedDocument:
    path: str
    sections: list = field(default_factory=list)
    format: str = ""
    note: str = ""

    @property
    def text(self) -> str:
        return "\n\n".join(s.text for s in self.sections)


class UnsupportedFormatError(ValueError):
    pass


def load_txt(path: str) -> LoadedDocument:
    with open(path, encoding="utf-8", errors="replace") as f:
        content = f.read()
    return LoadedDocument(path, [DocumentSection(content, source_file=path)], "txt")


def load_json(path: str) -> LoadedDocument:
    with open(path, encoding="utf-8") as f:
        data = json.load(f)

    sections = []

    def walk(node, trail="root"):
        if isinstance(node, dict):
            for k, v in node.items():
                walk(v, f"{trail}.{k}")
        elif isinstance(node, list):
            for i, v in enumerate(node):
                walk(v, f"{trail}[{i}]")
        elif node is not None and str(node).strip():
            sections.append(DocumentSection(f"{trail}: {node}", section=trail, source_file=path))

    walk(data)
    return LoadedDocument(path, sections, "json")


def load_csv(path: str) -> LoadedDocument:
    sections = []
    with open(path, encoding="utf-8", errors="replace", newline="") as f:
        reader = csv.DictReader(f)
        for i, row in enumerate(reader, start=1):
            parts = [f"{k}: {v}" for k, v in row.items() if v not in (None, "")]
            if parts:
                sections.append(
                    DocumentSection("; ".join(parts), section=f"row {i}", source_file=path)
                )
    return LoadedDocument(path, sections, "csv")


def load_pdf(path: str) -> LoadedDocument:
    try:
        from pypdf import PdfReader
    except ImportError:
        try:
            from PyPDF2 import PdfReader  # type: ignore
        except ImportError:
            raise UnsupportedFormatError(
                "PDF support requires pypdf. Install it with: pip install pypdf"
            )

    reader = PdfReader(path)
    sections = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            # page_number here is the real page index reported by the parser.
            sections.append(DocumentSection(text, page=page_number, source_file=path))
    note = "" if sections else "No extractable text (the PDF may be scanned images)."
    return LoadedDocument(path, sections, "pdf", note)


def load_docx(path: str) -> LoadedDocument:
    try:
        import docx  # python-docx
    except ImportError:
        raise UnsupportedFormatError(
            "DOCX support requires python-docx. Install it with: pip install python-docx"
        )

    document = docx.Document(path)
    sections = []
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            sections.append(DocumentSection(text, source_file=path))
    for t_i, table in enumerate(document.tables, start=1):
        for r_i, row in enumerate(table.rows, start=1):
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                sections.append(
                    DocumentSection(" | ".join(cells),
                                     section=f"table {t_i} row {r_i}", source_file=path)
                )
    return LoadedDocument(path, sections, "docx")


LOADERS = {
    ".txt": load_txt, ".md": load_txt, ".text": load_txt,
    ".json": load_json,
    ".csv": load_csv,
    ".pdf": load_pdf,
    ".docx": load_docx,
}


def load_document(path: str) -> LoadedDocument:
    if not os.path.exists(path):
        raise FileNotFoundError(path)
    ext = os.path.splitext(path)[1].lower()
    if ext not in LOADERS:
        raise UnsupportedFormatError(
            f"Unsupported format '{ext}'. Supported: {sorted(LOADERS)}"
        )
    return LOADERS[ext](path)
